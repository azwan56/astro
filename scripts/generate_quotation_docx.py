import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def create_quotation_docx(output_path):
    doc = Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("人类图数字化平台开发与技术服务报价及合作方案书")
        hrun.font.name = "PingFang SC"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(156, 163, 175)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("机密 · 仅供商务合作使用")
        frun.font.name = "PingFang SC"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(156, 163, 175)

    # Base font setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'PingFang SC'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(31, 41, 55)
    style_normal.paragraph_format.line_spacing = 1.35
    style_normal.paragraph_format.space_after = Pt(4)

    # --- Title Banner ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("人类图（Human Design）数字化平台")
    run_title.font.name = "PingFang SC"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(30, 58, 138) # Navy Blue

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("定制开发、系统上线与运维服务合作方案书（最终确认版）")
    run_sub.font.name = "PingFang SC"
    run_sub.font.size = Pt(13)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(75, 85, 99)

    # Metadata Info Box
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_table, color="E5E7EB", sz="4")
    meta_data = [
        [("客户方（甲方）：", "待指定企业"), ("服务方（乙方）：", "技术服务团队")],
        [("方案版本：", "V1.0 最终确认版"), ("编制日期：", "2026年8月")]
    ]
    for row_idx, row in enumerate(meta_data):
        for col_idx, (label, val) in enumerate(row):
            cell = meta_table.rows[row_idx].cells[col_idx]
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            r1 = cp.add_run(label)
            r1.font.bold = True
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = RGBColor(71, 85, 105)
            r2 = cp.add_run(val)
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Helper for Headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "PingFang SC"
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "PingFang SC"
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(14, 116, 144)
        return p

    # --- SECTION 1 ---
    add_heading_1("一、 商务费用汇总与核心条款")
    t1 = doc.add_table(rows=4, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1, color="CBD5E1", sz="6")
    
    # Headers
    headers1 = ["费用类别", "约定金额（人民币）", "服务范围与保障周期"]
    for i, h in enumerate(headers1):
        cell = t1.rows[0].cells[i]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 2 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    data1 = [
        ("项目定制开发总费用\n（商务优惠后）", "¥ 100,000 元\n(人民币壹拾万元整)", "包含人类图排盘算法引擎、科普知识库CMS、大师预约与咨询系统、用户端小程序/H5、运营管理后台开发、云端容器化部署及上线交付。"),
        ("上线第一年系统运维服务", "¥ 0 元\n(首年免费赠送)", "自系统正式验收上线之日起 365 天内，提供免费故障抢修、系统保活巡检、安全防护与程序 Bug 修复保障。"),
        ("次年起年度运维技术服务费", "¥ 10,000 元 / 年\n(每年人民币壹万元整)", "自上线满 1 年起按年计费（固定费用）。提供长期日常系统巡检、数据备份监督、微信官方规范升级适配及技术支持。")
    ]

    for row_idx, row in enumerate(data1, start=1):
        for col_idx, text in enumerate(row):
            cell = t1.rows[row_idx].cells[col_idx]
            bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.bold = True
                r.font.size = Pt(10)
            elif col_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.bold = True
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(185, 28, 28) if "100,000" in text or "10,000" in text else RGBColor(5, 150, 105)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(text)
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 2 ---
    add_heading_1("二、 双方权责划分（职责边界）")
    
    t2 = doc.add_table(rows=3, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2, color="CBD5E1", sz="6")
    headers2 = ["责任主体", "负责事项模块", "具体职责与交付内容"]
    for i, h in enumerate(headers2):
        cell = t2.rows[0].cells[i]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    data2 = [
        ("甲方\n（客户方）", "资质申请\n与资源提供", 
         "1. 提供企业有效营业执照及法人主体资质信息；\n"
         "2. 申请并持有微信公众平台/微信小程序企业认证、微信商户号（微信支付开通）；\n"
         "3. 采购并持有云服务器（推荐阿里云 ECS/FC）、域名、数据库等云资源；\n"
         "4. 配合完成域名 ICP 备案与公安联网备案；\n"
         "5. 提供人类图大师简介、资质证书、服务定价、科普文案及品牌视觉素材；\n"
         "6. 参与需求评审并对《产品需求规格说明书》(PRD) 进行最终签字确认。"),
        ("乙方\n（服务方）", "系统定制研发\n部署与运维", 
         "1. 梳理业务并输出完备的《产品需求规格说明书》(PRD) 与 UI 原型稿；\n"
         "2. 研发人类图底层星历计算引擎（0 外部付费 API 依赖，纯本地运行）；\n"
         "3. 研发用户端小程序/Web端、科普知识库、大师预约系统及全功能管理后台；\n"
         "4. 负责云端容器化环境搭建、系统部署上线、联调测试与小程序提审发布；\n"
         "5. 严格履行上线第一年免费运维保障，并自次年起提供有偿年度维保；\n"
         "6. 结清款项后交付完整系统源码、接口文档、数据库字典及后台操作手册。")
    ]

    for row_idx, row in enumerate(data2, start=1):
        for col_idx, text in enumerate(row):
            cell = t2.rows[row_idx].cells[col_idx]
            bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if col_idx < 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.bold = True
                r.font.size = Pt(10)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(text)
                r.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 3 ---
    add_heading_1("三、 项目开发模块与费用明细表（折后确认）")

    t3 = doc.add_table(rows=11, cols=5)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t3, color="CBD5E1", sz="6")
    headers3 = ["序号", "功能模块 / 服务内容", "详细工作内容与交付标准", "工期", "折后费用"]
    for i, h in enumerate(headers3):
        cell = t3.rows[0].cells[i]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=140, bottom=140, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    data3 = [
        ("1", "需求梳理与 UI 设计", "梳理全平台业务逻辑，输出高保真原型与全套 UI 界面设计稿，双方签字确认《产品需求规格说明书》(PRD)。", "7~8 天", "¥ 10,000"),
        ("2", "人类图核心算法引擎", "13天体黄经计算、88°太阳弧倒推、315°曼陀罗卦爻映射、9大中心/36通道/能量类型/内在权威判定、动态 BodyGraph 矢量 SVG 渲染（0外部API依赖）。", "已具备", "¥ 15,000"),
        ("3", "人类图理念科普中心\n(新增模块)", "前端分类浏览、理论体系起源（易经/脉轮/卡巴拉/占星）、四大类型图解、多媒体音视频讲解；后台富文本 CMS 内容发布与推荐位管理。", "4~5 天", "¥ 8,000"),
        ("4", "大师介绍与预约系统\n(新增模块)", "大师主页与资质展示、动态可约排班日历组件、生辰档案自动关联、微信在线支付预约、短信/微信服务号日程提醒通知、咨询评价；后台排班与调度台账。", "10~12 天", "¥ 22,000"),
        ("5", "用户端核心排盘应用\n(微信小程序 / H5)", "精准生辰输入与全球时区/夏令时校准、交互式人体图高亮、多盘管理（家人/朋友）、精美海报长图生成与微信分享、微信授权登录与绑定。", "10~12 天", "¥ 22,000"),
        ("6", "运营管理后台\n(Web 全功能版)", "用户档案管理、文案库管理、大师入驻与排班管理、预约订单流水与分成对账看板、系统权限配置与操作审计日志。", "8~10 天", "¥ 15,000"),
        ("7", "系统联调与安全测试", "接口联调、排盘精度校验、预约并发防超卖锁测试、微信支付闭环压测、系统性能调优与 Bug 清零。", "4~5 天", "¥ 5,000"),
        ("8", "云端部署上线与交付", "阿里云容器化（Docker）架构部署、SSL 证书配置、域名绑定、微信小程序提审发布、生产环境上线调优。", "3~4 天", "¥ 3,000"),
        ("9", "第一年系统技术运维", "上线首年免费赠送：7×24 小时应急响应、系统保活巡检、安全防护、缺陷 Bug 修复。", "365 天", "¥ 0 (赠送)"),
        ("—", "项目总费用与交付周期", "全套系统完整交付（含源代码、全套文档、部署上线、首年免费运维）。", "45~55工作日", "¥ 100,000 元")
    ]

    for row_idx, row in enumerate(data3, start=1):
        is_total = (row_idx == len(data3))
        for col_idx, text in enumerate(row):
            cell = t3.rows[row_idx].cells[col_idx]
            if is_total:
                bg = "FEF3C7" # Yellow Accent
            else:
                bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if col_idx in [0, 3, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.size = Pt(9) if not is_total else Pt(10)
            if is_total:
                r.font.bold = True
                if col_idx == 4:
                    r.font.color.rgb = RGBColor(185, 28, 28)
            elif col_idx == 4:
                r.font.bold = True
                r.font.color.rgb = RGBColor(185, 28, 28) if "¥ 0" not in text else RGBColor(5, 150, 105)
            elif col_idx == 1 and "新增" in text:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 4 ---
    add_heading_1("四、 阶段付款节点与实施里程碑")
    
    p_pay_intro = doc.add_paragraph("项目开发总款项为人民币 ¥ 100,000 元（壹拾万元整），按照项目实施里程碑分为 4 期进行结算：")
    p_pay_intro.runs[0].font.size = Pt(10)

    t4 = doc.add_table(rows=5, cols=4)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t4, color="CBD5E1", sz="6")
    headers4 = ["付款阶段", "款项比例与金额", "付款触发条件（里程碑标准）", "后续工作安排"]
    for i, h in enumerate(headers4):
        cell = t4.rows[0].cells[i]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=140, bottom=140, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    data4 = [
        ("第一期\n(启动款)", "40%\n¥ 40,000 元", "合同签署生效，双方完成需求对接并签字确认《产品需求规格说明书》(PRD) 后 3 个工作日内支付。", "启动 UI 设计、系统架构工程与核心模块开发。"),
        ("第二期\n(进度款)", "30%\n¥ 30,000 元", "完成排盘引擎、科普中心、小程序前台与大师预约核心流程，在测试环境提供功能演示并通过确认后 3 个工作日内支付。", "开展多端联调、微信支付接入、安全与并发压测。"),
        ("第三期\n(上线款)", "20%\n¥ 20,000 元", "完成生产环境部署，微信小程序审核通过并正式上线发布，乙方交付源码与全套文档后 3 个工作日内支付。", "系统正式对外运营，进入首年免费维保期。"),
        ("第四期\n(质保款)", "10%\n¥ 10,000 元", "系统正式上线稳定运行满 1~3 个月，无重大遗留缺陷后 3 个工作日内付清。", "完成最终质保闭环，持续提供第一年技术支持。")
    ]

    for row_idx, row in enumerate(data4, start=1):
        for col_idx, text in enumerate(row):
            cell = t4.rows[row_idx].cells[col_idx]
            bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if col_idx < 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.bold = True
                r.font.size = Pt(9.5)
                if col_idx == 1:
                    r.font.color.rgb = RGBColor(185, 28, 28)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(text)
                r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 5 ---
    add_heading_1("五、 需求基线与变更管理规则（CR 机制）")
    
    cr_points = [
        ("1. 需求基线锁定：", "双方共同梳理并签字确认的《产品需求规格说明书》（PRD）为本次项目实施的唯一技术基准与最终验收依据。系统严格按照签字文档所列的功能清单、业务逻辑及交互流程进行开发与测试验收。"),
        ("2. 需求变更定义：", "在项目执行过程中，凡超出签字确认版 PRD 范围的新增功能模块、业务流程大幅调整、新第三方系统对接或大幅界面重构，均属于需求变更（Change Request, CR）。"),
        ("3. 变更审批与计费流程：", "对于需求变更，乙方将在收到甲方书面变更说明后 2 个工作日内评估对架构、工期及费用的影响。超出基线范围的工作量统一按 人民币 ¥ 1,500 元 / 人天 进行核算。经双方书面确认追加费用及顺延工期并签署《项目补充协议》后，乙方统一安排排期执行。")
    ]
    for title, desc in cr_points:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(title)
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(15, 23, 42)
        r2 = p.add_run(desc)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(55, 65, 81)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 6 ---
    add_heading_1("六、 运维服务内容与 SLA 保障承诺")
    
    add_heading_2("1. 运维保障范围（首年免费 + 次年起每年固定 ¥ 10,000 元）")
    maint_items = [
        "系统缺陷修复（Bug Fix）：在合同 PRD 功能范围内的程序缺陷在维保期内提供免费排查与修复补丁；",
        "日常巡检与性能监控：定期检查云服务器 CPU、内存、磁盘存储及系统日志，确保服务持续稳定运行；",
        "数据灾备策略监督：协助配置数据库每日自动定时全量/增量备份，确保数据资产安全与可恢复性；",
        "安全与证书维护：协助进行 SSL 证书年度续期配置，修复基础运行环境已知的重大高危安全漏洞；",
        "微信平台兼容适配：针对微信小程序官方基础库升级与常规接口规则变动进行兼容性微调。"
    ]
    for item in maint_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.size = Pt(9.5)

    add_heading_2("2. SLA 故障响应与处置标准")
    t_sla = doc.add_table(rows=4, cols=4)
    t_sla.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_sla, color="CBD5E1", sz="6")
    headers_sla = ["故障级别", "故障定义与场景", "响应时效", "恢复与处置承诺"]
    for i, h in enumerate(headers_sla):
        cell = t_sla.rows[0].cells[i]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    sla_data = [
        ("P0 (特别重大)", "系统整体宕机、核心排盘算法瘫痪、预约微信支付通道完全阻断。", "30分钟内", "2小时内提供应急恢复方案或降级方案，7×24小时紧急抢修。"),
        ("P1 (主要故障)", "非核心功能异常（如科普图文显示错位、某位大师排班更新延迟），不影响主支付链路。", "2小时内", "8小时内定位并发布热修复补丁。"),
        ("P2 (轻微瑕疵)", "前端 UI 微小样式瑕疵、偶发非关键文案显示问题、常规后台操作疑问。", "4小时内", "1~2个工作日内随常规版本迭代修复。")
    ]
    for row_idx, row in enumerate(sla_data, start=1):
        for col_idx, text in enumerate(row):
            cell = t_sla.rows[row_idx].cells[col_idx]
            bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if col_idx in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.bold = True
                r.font.size = Pt(9)
                if "P0" in text:
                    r.font.color.rgb = RGBColor(185, 28, 28)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(text)
                r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 7 ---
    add_heading_1("七、 交付物与知识产权约定")
    deliv_points = [
        ("1. 知识产权与源码归属：", "在甲方按照约定结清全部开发款项后，本项目所产出的微信小程序前端代码、管理后台前端代码、业务后端服务源码、排盘算法模块及 Docker 部署脚本的完整源代码与知识产权完全归甲方所有。甲方拥有该软件系统的申请软件著作权、商业运营及收益的全部权利。"),
        ("2. 交付文档清单：", "包含《系统架构设计与数据字典说明书》、《RESTful API 接口规范文档》、《云服务器生产环境部署与运维手册》、《运营管理后台操作使用指南》以及《系统测试与验收报告》。")
    ]
    for title, desc in deliv_points:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(title)
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SECTION 8 ---
    add_heading_1("八、 双方确认与签署栏")
    
    t_sign = doc.add_table(rows=6, cols=2)
    t_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_sign, color="E5E7EB", sz="4")
    
    sign_headers = ["甲方（客户方）盖章签署", "乙方（服务方）盖章签署"]
    for i, h in enumerate(sign_headers):
        cell = t_sign.rows[0].cells[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(30, 58, 138)

    sign_rows = [
        ("企业名称（盖章）：", "企业名称（盖章）："),
        ("授权代表签字：", "授权代表签字："),
        ("联系电话：", "联系电话："),
        ("电子邮箱：", "电子邮箱："),
        ("签署日期：      年    月    日", "签署日期：      年    月    日")
    ]
    for row_idx, row in enumerate(sign_rows, start=1):
        for col_idx, text in enumerate(row):
            cell = t_sign.rows[row_idx].cells[col_idx]
            set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(55, 65, 81)

    # Save document
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    output_dir = "/Users/azwan/Projects/astro"
    output_file = os.path.join(output_dir, "人类图数字化平台开发与技术服务报价及合作方案书_最终确认版.docx")
    create_quotation_docx(output_file)
